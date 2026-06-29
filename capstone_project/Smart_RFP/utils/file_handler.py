"""
utils/file_handler.py
---------------------
Step 2 of the pipeline: extract raw text from an uploaded PDF / DOCX / TXT file
and clean it. Works on a file path or on raw bytes (Streamlit gives us bytes).
"""

import io
import re
from pypdf import PdfReader
import docx  # python-docx


def _clean(text: str) -> str:
    """Light cleaning: collapse whitespace, drop empty lines."""
    text = text.replace("\r", "\n")
    # collapse 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # collapse runs of spaces/tabs
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = [ln.strip() for ln in text.split("\n")]
    return "\n".join(ln for ln in lines if ln)


def extract_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def extract_from_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # also pull table cell text
    for table in document.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_from_txt(data: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_text(file_name: str, data: bytes) -> str:
    """
    Dispatch on file extension and return cleaned text.
    Raises ValueError for unsupported types.
    """
    name = (file_name or "").lower()
    if name.endswith(".pdf"):
        raw = extract_from_pdf(data)
    elif name.endswith(".docx"):
        raw = extract_from_docx(data)
    elif name.endswith(".doc"):
        # python-docx cannot read legacy .doc; ask user to convert
        raise ValueError("Legacy .doc is not supported. Please upload .docx, .pdf or .txt.")
    elif name.endswith(".txt"):
        raw = extract_from_txt(data)
    else:
        raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")
    cleaned = _clean(raw)
    if not cleaned.strip():
        raise ValueError("No readable text found in the file (it may be scanned or empty).")
    return cleaned
